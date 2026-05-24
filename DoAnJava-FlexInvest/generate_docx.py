from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_doc():
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN MÔN HỌC LẬP TRÌNH JAVA')
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LỚP IS216.xxx')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph('')
    
    # Details
    doc.add_paragraph('1. Tên nhóm: Nhóm ...')
    doc.add_paragraph('2. Tên đề tài: Ứng dụng Quản lý Tài chính và Đầu tư cá nhân FlexInvest')
    
    p = doc.add_paragraph('3. Tóm tắt nội dung đề tài (không quá 10 dòng):')
    p.runs[0].bold = True
    summary = ('FlexInvest là ứng dụng Desktop Java với mô hình MVC, cung cấp nền tảng quản lý tài chính toàn diện. '
              'Người dùng có thể quản lý số dư, nạp/rút tiền, gửi tiết kiệm, tham gia các gói đầu tư tài chính, '
              'thực hiện nhiệm vụ hàng ngày (check-in) để nhận thưởng, và nộp hồ sơ định danh điện tử (eKYC). '
              'Hệ thống tích hợp phân hệ Staff để nhân viên phê duyệt giao dịch/eKYC và phân hệ Admin để quản trị '
              'người dùng cùng cấu hình hệ thống. Ứng dụng sử dụng giao diện FlatLaf hiện đại, tích hợp JDBC '
              'với hệ quản trị cơ sở dữ liệu Oracle đảm bảo tính bảo mật và toàn vẹn dữ liệu.')
    doc.add_paragraph(summary)
    
    p = doc.add_paragraph('4. Thông tin nhóm:')
    p.runs[0].bold = True
    
    # Table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['STT', 'MSSV', 'Họ tên', 'Mô tả công việc được giao', 'Tổng công việc được phân (% công việc được phân công)', 'Khả năng hoàn thành công việc được giao (% hoàn thành)', 'Vai trò', 'Ký tên']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    members = [
        ('1', '21520001', 'Nguyễn Văn A', 'Thiết kế CSDL Oracle, xây dựng kiến trúc Model và DAO (BaseDAO), quản lý module Ví (Wallet), Giao dịch (Ledger), và xử lý luồng nạp/rút tiền backend.', '25%', '100%', 'Nhóm trưởng / Backend', ''),
        ('2', '21520002', 'Trần Thị B', 'Xây dựng giao diện Khách hàng (Dashboard, Wallet, Investments, Savings, Mission), xử lý sự kiện và tích hợp dữ liệu giao diện với DAO.', '25%', '100%', 'Thành viên / Frontend', ''),
        ('3', '21520003', 'Lê Văn C', 'Xây dựng luồng định danh điện tử (eKYC) và thiết kế toàn bộ phân hệ Staff (Duyệt eKYC, duyệt nạp/rút tiền) hỗ trợ nhân viên xử lý yêu cầu.', '25%', '100%', 'Thành viên / Fullstack', ''),
        ('4', '21520004', 'Phạm Thị D', 'Xây dựng module Xác thực (Đăng nhập, Đăng ký, Phân quyền), hệ thống Thông báo, phân hệ Admin, và tinh chỉnh UI/UX tổng thể (FlatLaf, Icons).', '25%', '100%', 'Thành viên / Fullstack', ''),
    ]
    
    for member in members:
        row_cells = table.add_row().cells
        for i, text in enumerate(member):
            row_cells[i].text = text
            if i in [0, 1, 4, 5]:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.save('PhanCongCongViec_FlexInvest.docx')
    print('Document saved.')

if __name__ == '__main__':
    create_doc()
